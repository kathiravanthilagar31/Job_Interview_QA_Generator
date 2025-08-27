from langchain.document_loaders import PyPDFLoader
from langchain.text_splitter import TokenTextSplitter
from langchain.docstore.document import Document
from langchain.chat_models import ChatOpenAI
from langchain.prompts import PromptTemplate
from langchain.chains.summarize import load_summarize_chain
from langchain.embeddings.openai import OpenAIEmbeddings
from langchain.vectorstores import FAISS
from langchain.chains import RetrievalQA
import os
from dotenv import load_dotenv
import docx  # Import python-docx
from src.prompt import *

# OpenAI API setup
load_dotenv()
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
os.environ["OPENAI_API_KEY"] = OPENAI_API_KEY

# Function to extract text from DOCX file
def extract_text_from_docx(docx_path):
    doc = docx.Document(docx_path)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    return '\n'.join(full_text)

# File processing and Chunking operation
def file_processing(file_path=None, jd_text=None):
    # Load data from file or JD text
    if file_path:
        if file_path.endswith('.pdf'):
            loader = PyPDFLoader(file_path)
            data = loader.load()
        elif file_path.endswith('.docx'):
            text = extract_text_from_docx(file_path)  # Use python-docx to extract text
            data = text.split('\n')  # Split the text into lines for processing
        else:
            raise ValueError("Unsupported file type. Please provide a PDF or DOCX file.")
    elif jd_text:
        data = jd_text.split('\n')  # Split pasted JD text by new lines for processing
    else:
        raise ValueError("Either a file path or JD text must be provided.")
    
    # Convert documents to string
    question = ""
    for page in data:
        question += page if isinstance(page, str) else page.page_content  # Handle docx page content
    
    # Chunking operations for Question Generation
    splitter_question = TokenTextSplitter(
        model_name="gpt-3.5-turbo",
        chunk_size=1000,
        chunk_overlap=100
    )
    chunk_question_gen = splitter_question.split_text(question)
    
    # Convert chunks to documents for question generation
    doc_ques_gen = [Document(page_content=t) for t in chunk_question_gen]
    
    # Chunking operations for Answer Generation
    splitter_ans_gen = TokenTextSplitter(
        model_name="gpt-3.5-turbo",
        chunk_size=1000,
        chunk_overlap=100
    )
    doc_ans_gen = splitter_ans_gen.split_documents(doc_ques_gen)
    
    return doc_ques_gen, doc_ans_gen

# LLM Pipeline and Vector Embeddings
def llm_pipeline(file_path=None, jd_text=None):
    doc_ques_gen, doc_ans_gen = file_processing(file_path, jd_text)
    
    # LLM for Question Generation
    llm_ques_gen_pipeline = ChatOpenAI(
        model="gpt-3.5-turbo",
        temperature=0.3
    )
    
    # Prompt templates for question generation
    PROMPT_QUESTIONS = PromptTemplate(
        template=prompt_template,
        input_variables=['text']  # user's uploading document or pasted JD text
    )
    
    REFINE_PROMPT_QUESTIONS = PromptTemplate(
        input_variables=['existing_answer', "text"],
        template=refine_template
    )
    
    ques_gen_chain = load_summarize_chain(
        llm=llm_ques_gen_pipeline,
        chain_type="refine",
        question_prompt=PROMPT_QUESTIONS,
        refine_prompt=REFINE_PROMPT_QUESTIONS
    )
    
    questions = ques_gen_chain.run(doc_ques_gen)
    questions_list = questions.split('\n')
    
    # Embeddings and VectorStore for Answer Generation
    embeddings = OpenAIEmbeddings()
    vector_store = FAISS.from_documents(doc_ans_gen, embeddings)
    
    llm_ans_gen = ChatOpenAI(temperature=0.1, model="gpt-3.5-turbo")
    
    answer_generation_chain = RetrievalQA.from_chain_type(
        llm=llm_ans_gen,
        chain_type="stuff",
        retriever=vector_store.as_retriever()
    )
    
    return answer_generation_chain, questions_list
    